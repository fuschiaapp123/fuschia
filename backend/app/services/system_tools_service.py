"""
System Tools Service for providing RAG and MCP-powered tools to workflow agents.

This service provides a framework for extensible system tools that can:
1. Access knowledge bases through RAG (Retrieval Augmented Generation)
2. Integrate with MCP (Model Context Protocol) services
3. Provide context-aware assistance to workflow execution agents
"""

import asyncio
import json
import uuid
import boto3
from typing import Dict, List, Optional, Any, Callable
from datetime import datetime
from abc import ABC, abstractmethod
import structlog
from dataclasses import dataclass
from enum import Enum
import os
from pathlib import Path

import PyPDF2
import docx
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings
from langchain_community.vectorstores import FAISS
from langchain.schema import Document


logger = structlog.get_logger()


class SystemToolCategory(Enum):
    RAG_KNOWLEDGE = "rag_knowledge"
    MCP_INTEGRATION = "mcp_integration"  
    CONTEXT_ENHANCEMENT = "context_enhancement"
    DATA_ACCESS = "data_access"
    EXTERNAL_SERVICES = "external_services"


@dataclass
class SystemToolMetadata:
    name: str
    category: SystemToolCategory
    description: str
    version: str
    requires_auth: bool = False
    async_capable: bool = True
    dependencies: List[str] = None
    
    def __post_init__(self):
        if self.dependencies is None:
            self.dependencies = []


class BaseSystemTool(ABC):
    """Base class for all system tools"""
    
    def __init__(self, metadata: SystemToolMetadata):
        self.metadata = metadata
        self.logger = logger.bind(tool=metadata.name)
        self.initialized = False
    
    @abstractmethod
    async def initialize(self) -> bool:
        """Initialize the tool and its dependencies"""
        pass
    
    @abstractmethod
    async def execute(self, *args, **kwargs) -> str:
        """Execute the tool's main functionality"""
        pass
    
    async def cleanup(self):
        """Clean up resources when tool is no longer needed"""
        pass
    
    def get_dspy_function(self) -> Callable:
        """Get a DSPy-compatible function for this tool - subclasses should override for specific parameters"""
        async def tool_wrapper(*args, **kwargs):
            if not self.initialized:
                await self.initialize()
            return await self.execute(*args, **kwargs)
        
        tool_wrapper.__name__ = self.metadata.name
        tool_wrapper.__doc__ = self.metadata.description
        return tool_wrapper


class RAGKnowledgeTool(BaseSystemTool):
    """RAG-powered tool for accessing knowledge bases stored in AWS S3 gather more information about a task"""
    
    def __init__(self):
        super().__init__(SystemToolMetadata(
            name="rag_knowledge_search",
            category=SystemToolCategory.RAG_KNOWLEDGE,
            description="Search through knowledge base of documents stored in AWS S3 using RAG",
            version="1.0.0",
            requires_auth=True,
            dependencies=["boto3", "langchain", "faiss", "openai"]
        ))
        logger.info("RAGKnowledgeTool initialized")
        # Configuration
        self.s3_bucket = os.getenv('KNOWLEDGE_S3_BUCKET', 'fuschia-knowledge-base')
        self.s3_prefix = os.getenv('KNOWLEDGE_S3_PREFIX', 'documents/')
        self.embedding_model = None
        self.vector_store = None
        self.s3_client = None
        self.local_cache_dir = Path("./knowledge_cache")
    
    async def initialize(self) -> bool:
        """Initialize S3 connection and embeddings"""
        try:
            # Quick configuration checks to fail fast
            
            
            openai_api_key = os.getenv('OPENAI_API_KEY')
            if not openai_api_key:
                self.logger.error("OpenAI API key not found for embeddings")
                return False
            
            aws_access_key = os.getenv('AWS_ACCESS_KEY_ID')
            aws_secret_key = os.getenv('AWS_SECRET_ACCESS_KEY')
            if not aws_access_key or not aws_secret_key:
                self.logger.error("AWS credentials not found for S3 access")
                return False
            
            # Initialize S3 client
            self.s3_client = boto3.client('s3',
                aws_access_key_id=aws_access_key,
                aws_secret_access_key=aws_secret_key,
                region_name=os.getenv('AWS_REGION', 'us-east-1')
            )
            
            # Test S3 connectivity quickly
            try:
                self.s3_client.head_bucket(Bucket=self.s3_bucket)
                self.logger.info("S3 bucket accessible", bucket=self.s3_bucket)
            except Exception as e:
                self.logger.warning("S3 bucket not accessible, continuing without S3", bucket=self.s3_bucket, error=str(e))
                # Continue without S3 - we can still provide fallback responses
            
            # Initialize embeddings
            try:
                self.embedding_model = OpenAIEmbeddings(openai_api_key=openai_api_key)
                self.logger.info("OpenAI embeddings model initialized successfully")
            except Exception as e:
                self.logger.error("Failed to initialize OpenAI embeddings", error=str(e))
                return False
            
            # Create local cache directory
            self.local_cache_dir.mkdir(exist_ok=True)
            
            # Initialize or load vector store (with timeout)
            try:
                await asyncio.wait_for(self._initialize_vector_store(), timeout=20.0)
            except asyncio.TimeoutError:
                self.logger.warning("Vector store initialization timed out, continuing with empty store")
                # Continue with empty vector store - tool will return appropriate message
            except Exception as e:
                self.logger.warning("Vector store initialization failed, continuing with empty store", error=str(e))
            
            self.initialized = True
            self.logger.info("RAG Knowledge Tool initialized successfully")
            return True
            
        except Exception as e:
            self.logger.error("Failed to initialize RAG Knowledge Tool", error=str(e))
            return False
    
    async def _initialize_vector_store(self):
        """Initialize or load the FAISS vector store"""
        vector_store_path = self.local_cache_dir / "vector_store"
        
        if vector_store_path.exists():
            try:
                # Load existing vector store
                self.vector_store = FAISS.load_local(str(vector_store_path), self.embedding_model)
                self.logger.info("Loaded existing vector store")
                return
            except Exception as e:
                self.logger.warning("Failed to load existing vector store, rebuilding", error=str(e))
        
        # Build new vector store from S3 documents
        await self._rebuild_vector_store()
    
    async def _rebuild_vector_store(self):
        """Rebuild vector store from documents in S3"""
        try:
            # List documents in S3 bucket
            response = self.s3_client.list_objects_v2(
                Bucket=self.s3_bucket,
                Prefix=self.s3_prefix
            )
            logger.info("Fetched document list from S3", document_count=len(response.get('Contents', [])))
            documents = []
            for obj in response.get('Contents', []):
                if obj['Key'].lower().endswith(('.pdf', '.docx', '.txt')):
                    doc_content = await self._process_s3_document(obj['Key'])
                    if doc_content:
                        documents.extend(doc_content)
            logger.info("Processed documents from S3", processed_count=len(documents))
            if documents:
                # Check if embedding model is available
                if not self.embedding_model:
                    self.logger.error("Embedding model not initialized - cannot create vector store")
                    return
                    
                # Create vector store
                logger.info("Building vector store from documents")
                self.vector_store = FAISS.from_documents(documents, self.embedding_model)
                
                # Validate vector store was created
                if self.vector_store and hasattr(self.vector_store, 'index'):
                    # Get vector count using FAISS index.ntotal property
                    vector_count = getattr(self.vector_store.index, 'ntotal', 0)
                    logger.info("Vector store built successfully", vector_count=vector_count)
                else:
                    logger.error("Vector store creation failed - invalid vector store object")
                    return

                # Save vector store locally
                
                vector_store_path = self.local_cache_dir / "vector_store"
                logger.info("Saving vector store locally", path=str(vector_store_path))
                self.vector_store.save_local(str(vector_store_path))
                logger.info("Saved vector store locally", path=str(vector_store_path))  
                
                self.logger.info(f"Built vector store with {len(documents)} document chunks")
            else:
                self.logger.warning("No documents found in S3 bucket")
                
        except Exception as e:
            self.logger.error("Failed to rebuild vector store", error=str(e))
    
    async def _process_s3_document(self, s3_key: str) -> List[Document]:
        """Download and process a document from S3"""
        try:
            # Download document from S3
            local_path = self.local_cache_dir / s3_key.replace('/', '_')
            self.s3_client.download_file(self.s3_bucket, s3_key, str(local_path))
            
            # Extract text based on file type
            text_content = ""
            file_extension = local_path.suffix.lower()
            
            if file_extension == '.pdf':
                text_content = self._extract_pdf_text(local_path)
            elif file_extension == '.docx':
                text_content = self._extract_docx_text(local_path)
            elif file_extension == '.txt':
                with open(local_path, 'r', encoding='utf-8') as f:
                    text_content = f.read()
            
            if not text_content.strip():
                return []
            
            # Split text into chunks
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=1000,
                chunk_overlap=200,
                separators=["\n\n", "\n", " ", ""]
            )
            
            chunks = text_splitter.split_text(text_content)
            
            # Create Document objects
            documents = []
            for i, chunk in enumerate(chunks):
                doc = Document(
                    page_content=chunk,
                    metadata={
                        "source": s3_key,
                        "chunk_id": i,
                        "file_type": file_extension,
                        "processed_at": datetime.utcnow().isoformat()
                    }
                )
                documents.append(doc)
            
            # Clean up local file
            local_path.unlink(missing_ok=True)
            
            return documents
            
        except Exception as e:
            self.logger.error(f"Failed to process document {s3_key}", error=str(e))
            return []
    
    def _extract_pdf_text(self, file_path: Path) -> str:
        """Extract text from PDF file"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text()
                return text
        except Exception as e:
            self.logger.error(f"Failed to extract PDF text from {file_path}", error=str(e))
            return ""
    
    def _extract_docx_text(self, file_path: Path) -> str:
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            self.logger.error(f"Failed to extract DOCX text from {file_path}", error=str(e))
            return ""
    
    async def execute(self, query: str, max_results: int = 5) -> str:
        """Search the knowledge base for gathering information on how to resolve a task"""
        try:
            if not self.vector_store:
                return "Knowledge base not available. Please check the configuration."
            logger.info("Executing RAG knowledge search", query=query)  
            # Perform similarity search
            results = self.vector_store.similarity_search(query, k=max_results)
            logger.info("RAG search results retrieved", result_count=len(results))
            if not results:
                return f"No relevant information found for query: {query}"
            
            # Format results
            formatted_results = []
            for i, doc in enumerate(results, 1):
                source = doc.metadata.get('source', 'Unknown')
                chunk_id = doc.metadata.get('chunk_id', 0)
                content = doc.page_content.strip()
                
                formatted_results.append(
                    f"**Result {i}** (Source: {source}, Chunk: {chunk_id}):\n{content}\n"
                )
            
            response = f"Found {len(results)} relevant documents for query '{query}':\n\n"
            response += "\n".join(formatted_results)
            
            self.logger.info("RAG search completed", query=query, results_count=len(results))
            return response
            
        except Exception as e:
            error_msg = f"RAG search failed: {str(e)}"
            self.logger.error("RAG search execution failed", query=query, error=str(e))
            return error_msg
    
    def get_dspy_function(self) -> Callable:
        """Get a DSPy-compatible function with proper parameter definitions"""
        async def rag_knowledge_search(query: str, max_results: int = 5) -> str:
            """Search through knowledge base of documents stored in AWS S3 using RAG"""
            if not self.initialized:
                try:
                    # Try to initialize with a shorter timeout
                    init_success = await asyncio.wait_for(self.initialize(), timeout=30.0)
                    if not init_success:
                        return f"RAG knowledge base is not available. Configuration issues detected. Query: {query}"
                except asyncio.TimeoutError:
                    self.logger.warning("RAG tool initialization timed out", query=query)
                    return f"RAG knowledge base initialization timed out. Unable to search for: {query}"
                except Exception as e:
                    self.logger.error("RAG tool initialization failed", query=query, error=str(e))
                    return f"RAG knowledge base is unavailable due to configuration issues. Cannot search for: {query}"
            
            try:
                # Execute with timeout to prevent hanging
                result = await asyncio.wait_for(self.execute(query=query, max_results=max_results), timeout=25.0)
                return result
            except asyncio.TimeoutError:
                self.logger.warning("RAG search execution timed out", query=query)
                return f"RAG search timed out for query: {query}. Try a more specific search term."
            except Exception as e:
                self.logger.error("RAG search execution failed", query=query, error=str(e))
                return f"RAG search failed for query: {query}. Error: {str(e)}"
        
        rag_knowledge_search.__name__ = self.metadata.name
        rag_knowledge_search.__doc__ = self.metadata.description
        return rag_knowledge_search


class MCPIntegrationTool(BaseSystemTool):
    """Tool for integrating with Model Context Protocol (MCP) services"""
    
    def __init__(self):
        super().__init__(SystemToolMetadata(
            name="mcp_service_call",
            category=SystemToolCategory.MCP_INTEGRATION,
            description="Call external MCP services to enhance context and capabilities",
            version="1.0.0",
            requires_auth=False,
            dependencies=["httpx"]
        ))
        
        # Instead of HTTP endpoints, directly use the MCP servers within the application
        self.mcp_servers = {}
    
    async def initialize(self) -> bool:
        """Initialize MCP connections by directly importing available MCP servers"""
        try:
            self.available_services = {}

            # Import and register available MCP servers
            try:
                from app.services.gmail_mcp_server import gmail_mcp_server
                if gmail_mcp_server.is_running:
                    self.mcp_servers["gmail"] = gmail_mcp_server
                    self.available_services["gmail"] = "direct"
                    self.logger.info("MCP Gmail service connected (direct)")
                else:
                    self.logger.debug("Gmail MCP server not running")
            except ImportError:
                self.logger.debug("Gmail MCP server not available")

            try:
                from app.services.hcmpro_mcp_server import hcmpro_mcp_server
                if hcmpro_mcp_server.is_running:
                    self.mcp_servers["hcmpro"] = hcmpro_mcp_server
                    self.available_services["hcmpro"] = "direct"
                    self.logger.info("MCP HCMPro service connected (direct)")
                else:
                    self.logger.debug("HCMPro MCP server not running")
            except ImportError:
                self.logger.debug("HCMPro MCP server not available")

            # Initialize even if no services are available (graceful degradation)
            self.initialized = True
            if self.available_services:
                self.logger.info(f"MCP Integration Tool initialized with {len(self.available_services)} direct services", services=list(self.available_services.keys()))
            else:
                self.logger.info("MCP Integration Tool initialized with no direct services (will use fallback)")
            return True

        except Exception as e:
            self.logger.error("Failed to initialize MCP Integration Tool", error=str(e))
            return False
    
    async def execute(self, service: str, method: str, parameters: Dict[str, Any] = None) -> str:
        """Call an MCP service with given method and parameters"""
        try:
            if not hasattr(self, 'mcp_servers'):
                return "MCP Integration Tool not properly initialized"
            service = service.split('_')[0]  # Extract base service name
            service = service.lower()
            

            # Check if the service is actually available
            if not hasattr(self, 'available_services') or service not in self.available_services:
                # Provide fallback for knowledge service
                if service == "knowledge" and method == "search":
                    return await self._fallback_knowledge_search(parameters or {})
                else:
                    available_services = list(getattr(self, 'available_services', {}).keys())
                    return f"MCP {service} service is not available. Available services: {available_services}"

            # Get the direct MCP server instance
            mcp_server = self.mcp_servers.get(service)
            if not mcp_server:
                return f"MCP {service} server instance not found"

            if parameters is None:
                parameters = {}
            self.logger.debug("MCP service call", service=service, method=method, parameters=parameters)    

            # Call the MCP server directly based on the method
            if method == "search" and service == "gmail":
                # Gmail search functionality
                result = await mcp_server.call_tool("gmail_list_messages", parameters)
                return json.dumps(result, indent=2)
            elif method == "send_message" and service == "gmail":
                # Gmail send functionality
                result = await mcp_server.call_tool("gmail_send_message", parameters)
                return json.dumps(result, indent=2)
            elif method == "list_job_offers" and service == "hcmpro":
                # HCMPro list job offers functionality
                self.logger.debug("Calling HCMPro list_job_offers", parameters=parameters)
                result = await mcp_server.call_tool("hcmpro_list_job_offers", parameters)
                return json.dumps(result, indent=2)
            elif method == "get_job_offer_record" and service == "hcmpro":
                # HCMPro search job offer by candidate name/email
                self.logger.debug("Calling HCMPro search_job_offers_by_candidate", parameters=parameters)
                # Map employee_name to candidate_name for backward compatibility
                if "employee_name" in parameters:
                    parameters["candidate_name"] = parameters.pop("employee_name")
                result = await mcp_server.call_tool("hcmpro_search_job_offers_by_candidate", parameters)
                return json.dumps(result, indent=2)
            else:
                # Generic tool call - try to map method to tool name
                tool_name = f"{service}_{method}"
                if hasattr(mcp_server, 'call_tool'):
                    result = await mcp_server.call_tool(tool_name, parameters)
                    return json.dumps(result, indent=2)
                else:
                    return f"Method '{method}' not supported for service '{service}'"

        except Exception as e:
            error_msg = f"MCP service call failed: {str(e)}"
            self.logger.error("MCP execution failed", service=service, method=method, error=str(e))
            return error_msg

    async def _fallback_knowledge_search(self, parameters: Dict[str, Any]) -> str:
        """Fallback knowledge search using the existing RAG system"""
        try:
            query = parameters.get("query", "")
            if not query:
                return "Error: No query provided for knowledge search"

            # Use the existing RAG knowledge search tool
            from app.services.system_tools_service import system_tools_service
            rag_tool = system_tools_service.get_tool("system_rag_knowledge_search")

            if rag_tool:
                result = await rag_tool.execute(query)
                return f"Knowledge search result (via RAG fallback): {result}"
            else:
                return "Error: Neither MCP knowledge service nor RAG knowledge search is available"

        except Exception as e:
            self.logger.error("Fallback knowledge search failed", error=str(e))
            return f"Fallback knowledge search failed: {str(e)}"
    
    def get_dspy_function(self) -> Callable:
        """Get a DSPy-compatible function with proper parameter definitions"""
        async def mcp_service_call(service: str, method: str, parameters: str = "{}") -> str:
            """Call external MCP services to enhance context and capabilities"""
            import json
            if not self.initialized:
                await self.initialize()
            
            # Parse parameters JSON string
            try:
                params_dict = json.loads(parameters) if parameters else {}
            except json.JSONDecodeError:
                params_dict = {}
            
            return await self.execute(service=service, method=method, parameters=params_dict)
        
        mcp_service_call.__name__ = self.metadata.name
        mcp_service_call.__doc__ = self.metadata.description
        return mcp_service_call


class ContextEnhancementTool(BaseSystemTool):
    """Tool for enhancing task context with additional information"""
    
    def __init__(self):
        super().__init__(SystemToolMetadata(
            name="enhance_context",
            category=SystemToolCategory.CONTEXT_ENHANCEMENT,
            description="Enhance task context with additional relevant information",
            version="1.0.0",
            requires_auth=False
        ))
    
    async def initialize(self) -> bool:
        self.initialized = True
        return True
    
    async def execute(self, task_description: str, current_context: Dict[str, Any] = None) -> str:
        """Enhance context for a given task"""
        try:
            if current_context is None:
                current_context = {}
            
            # Analyze task to determine what additional context might be helpful
            enhanced_context = {
                "timestamp": datetime.utcnow().isoformat(),
                "task_analysis": {
                    "description": task_description,
                    "complexity": self._assess_complexity(task_description),
                    "estimated_duration": self._estimate_duration(task_description),
                    "required_resources": self._identify_resources(task_description)
                },
                "suggestions": self._generate_suggestions(task_description),
                "original_context": current_context
            }
            
            return json.dumps(enhanced_context, indent=2)
            
        except Exception as e:
            error_msg = f"Context enhancement failed: {str(e)}"
            self.logger.error("Context enhancement failed", error=str(e))
            return error_msg
    
    def get_dspy_function(self) -> Callable:
        """Get a DSPy-compatible function with proper parameter definitions"""
        async def enhance_context(task_description: str, current_context: str = "{}") -> str:
            """Enhance task context with additional relevant information"""
            import json
            if not self.initialized:
                await self.initialize()
            
            # Parse context JSON string
            try:
                context_dict = json.loads(current_context) if current_context else {}
            except json.JSONDecodeError:
                context_dict = {}
            
            return await self.execute(task_description=task_description, current_context=context_dict)
        
        enhance_context.__name__ = self.metadata.name
        enhance_context.__doc__ = self.metadata.description
        return enhance_context
    
    def _assess_complexity(self, description: str) -> str:
        """Assess task complexity based on description"""
        keywords_high = ["integrate", "implement", "develop", "design", "architect"]
        keywords_medium = ["configure", "setup", "modify", "update", "analyze"]
        keywords_low = ["check", "verify", "list", "show", "display"]
        
        desc_lower = description.lower()
        
        if any(keyword in desc_lower for keyword in keywords_high):
            return "high"
        elif any(keyword in desc_lower for keyword in keywords_medium):
            return "medium"
        elif any(keyword in desc_lower for keyword in keywords_low):
            return "low"
        else:
            return "medium"  # default
    
    def _estimate_duration(self, description: str) -> str:
        """Estimate task duration based on complexity"""
        complexity = self._assess_complexity(description)
        duration_map = {
            "low": "5-15 minutes",
            "medium": "30-60 minutes", 
            "high": "2-4 hours"
        }
        return duration_map.get(complexity, "30-60 minutes")
    
    def _identify_resources(self, description: str) -> List[str]:
        """Identify resources that might be needed"""
        resources = []
        desc_lower = description.lower()
        
        if "database" in desc_lower or "sql" in desc_lower:
            resources.append("database_access")
        if "api" in desc_lower or "service" in desc_lower:
            resources.append("api_credentials")
        if "file" in desc_lower or "document" in desc_lower:
            resources.append("file_system_access")
        if "email" in desc_lower or "notification" in desc_lower:
            resources.append("notification_service")
        
        return resources
    
    def _generate_suggestions(self, description: str) -> List[str]:
        """Generate helpful suggestions for the task"""
        suggestions = []
        desc_lower = description.lower()
        
        if "error" in desc_lower or "bug" in desc_lower:
            suggestions.append("Check logs for error details")
            suggestions.append("Verify system configuration")
        
        if "performance" in desc_lower:
            suggestions.append("Monitor resource usage during execution")
            suggestions.append("Consider caching strategies")
        
        if "security" in desc_lower:
            suggestions.append("Review security best practices")
            suggestions.append("Ensure proper authentication")
        
        if not suggestions:
            suggestions.append("Break down complex tasks into smaller steps")
            suggestions.append("Document any assumptions or constraints")
        
        return suggestions


class SystemToolsService:
    """Service for managing and providing system tools to workflow agents"""
    
    def __init__(self):
        self.logger = logger.bind(service="SystemToolsService")
        self.tools: Dict[str, BaseSystemTool] = {}
        self.initialized = False
    
    async def initialize(self):
        """Initialize all system tools"""
        try:
            # Register default system tools
            await self._register_default_tools()
            
            # Initialize all tools
            for tool_name, tool in self.tools.items():
                try:
                    success = await tool.initialize()
                    if success:
                        self.logger.info("System tool initialized", tool=tool_name)
                    else:
                        self.logger.warning("System tool failed to initialize", tool=tool_name)
                except Exception as e:
                    self.logger.error("System tool initialization error", tool=tool_name, error=str(e))
            
            self.initialized = True
            self.logger.info("System Tools Service initialized", tool_count=len(self.tools))
            
        except Exception as e:
            self.logger.error("Failed to initialize System Tools Service", error=str(e))
            raise
    
    async def _register_default_tools(self):
        """Register default system tools"""
        # RAG Knowledge Tool
        rag_tool = RAGKnowledgeTool()
        await self.register_tool(rag_tool)
        
        # MCP Integration Tool  
        mcp_tool = MCPIntegrationTool()
        await self.register_tool(mcp_tool)
        
        # Context Enhancement Tool
        context_tool = ContextEnhancementTool()
        await self.register_tool(context_tool)
    
    async def register_tool(self, tool: BaseSystemTool):
        """Register a system tool"""
        self.tools[tool.metadata.name] = tool
        self.logger.info("Registered system tool", tool=tool.metadata.name, category=tool.metadata.category.value)
    
    def get_tool(self, tool_name: str) -> Optional[BaseSystemTool]:
        """Get a system tool by name"""
        return self.tools.get(tool_name)
    
    def get_tools_by_category(self, category: SystemToolCategory) -> List[BaseSystemTool]:
        """Get all tools in a specific category"""
        return [tool for tool in self.tools.values() if tool.metadata.category == category]
    
    def get_all_tools(self) -> Dict[str, BaseSystemTool]:
        """Get all registered system tools"""
        return self.tools.copy()
    
    def get_dspy_tools(self) -> List[Callable]:
        """Get DSPy-compatible functions for all tools (including those with graceful error handling)"""
        dspy_tools = []
        
        for tool_name, tool in self.tools.items():
            # Include all tools - they handle their own initialization failures gracefully
            dspy_func = tool.get_dspy_function()
            dspy_tools.append(dspy_func)
            
            if tool.initialized:
                self.logger.debug(f"DSPy tool ready: {tool_name}")
            else:
                self.logger.debug(f"DSPy tool with graceful degradation: {tool_name}")
        
        return dspy_tools
    
    async def cleanup(self):
        """Cleanup all system tools"""
        for tool in self.tools.values():
            try:
                await tool.cleanup()
            except Exception as e:
                self.logger.error("Error cleaning up tool", tool=tool.metadata.name, error=str(e))


# Global system tools service instance
system_tools_service = SystemToolsService()